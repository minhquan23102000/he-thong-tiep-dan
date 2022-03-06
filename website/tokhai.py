import ast
import io
import json

import docx
import requests
from chatbot.models import *
from flask import (Blueprint, Response, flash, jsonify, render_template,
                   request, session)
from flask_login import current_user

from website import db

from . import cloud_storage, dao, secret

TOKHAI_USER_PATH = 'tokhai_user'
TOKHAI_MAU_PATH = 'tokhai_mau'

tokhai = Blueprint("tokhai", __name__)

@tokhai.route('/dang-ky-ket-hon')
def dang_ky_ket_hon():
    return render_template("tokhai_kethon.html")


@tokhai.route('/dien-to-khai-tu-dong')
def dien_to_khai_tu_dong():
    
    paper_name = request.args.get('paper_name').strip()
    request_data = request.args.get('data').strip()
    request_data = ast.literal_eval(request_data)
    
    data = make_data_dien_tokhai(paper_name, request_data)
    
    buffer = dien_to_khai(paper_name, data)
    
    #Upload to cloud and get url
    if session.get('conversation_id') == None:
        conversation = dao.new_conversation()
        session['conversation_id'] = conversation.id
        conversation_id = conversation.id
    else:
        conversation_id = session.get('conversation_id')
    
    url = cloud_storage.upload_file(buffer=buffer, cloud_path=TOKHAI_USER_PATH, file_name=f'{paper_name}_{conversation_id}.docx')
    
    send_data = {'url': url}
    
    return jsonify(send_data)


def make_data_dien_tokhai(paper_name:str, data:dict):
    rs = {}
    if paper_name == 'dang_ky_ket_hon':
        rs = data.copy()
        rs['nam_tuythan'] = data['nam_loaigiaytuythan'] + " số " + data['nam_id']
        rs['nu_tuythan'] = data['nu_loaigiaytuythan'] + " số " + data['nam_id']
        
        rs.pop('nam_id')
        rs.pop('nu_id')
        rs.pop('nam_loaigiaytuythan')
        rs.pop('nu_loaigiaytuythan')

    return rs


def dien_to_khai(paper_name:str, data:dict):
    
    paper_path = f'website/static/tokhai_mau/{paper_name}.docx'
    template_document = docx.Document(paper_path) 
    
    for variable_key, variable_value in data.items():
        for paragraph in template_document.paragraphs:
            paragraph.text = paragraph.text.replace(variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.text = paragraph.text.replace(variable_key, variable_value)
    # save document info
    buffer = io.BytesIO()
    template_document.save(buffer)  # save your memory stream
    buffer.seek(0)  # rewind the stream                 
    return buffer

    
def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)

